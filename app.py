import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="Zodion - Auditoría Técnica", layout="wide")

st.title("🛡️ Sistema Profesional de Diagnóstico ZODION")

# SIDEBAR
with st.sidebar:
    cliente = st.text_input("Cliente", "JAVERIANO")
    fecha = st.date_input("Fecha", datetime.now())
    auditor = st.text_input("Auditor", "Supervisor Ambiental Zodion")

# TAB
tab1, tab2, tab3 = st.tabs(["📸 Evidencia", "🔍 Evaluación", "🧠 Diagnóstico"])

# -------------------------
# 1. EVIDENCIA FOTOGRÁFICA
# -------------------------
analisis_fotos = []

with tab1:
    fotos = st.file_uploader("Cargar fotos", type=["jpg","png"], accept_multiple_files=True)

    if fotos:
        for i, f in enumerate(fotos):
            st.image(f, width=250)
            titulo = st.text_input(f"Título {i+1}", key=f"t{i}")
            desc = st.text_area(f"Análisis técnico {i+1}", key=f"d{i}")

            # Mejora automática del análisis
            analisis_pro = f"""
EVIDENCIA {i+1}: {titulo.upper()}

Análisis técnico:
{desc}

Interpretación profesional:
Se identifican posibles factores de riesgo asociados a contaminación cruzada,
deficiencias en almacenamiento o fallas en condiciones higiénico-sanitarias,
lo cual puede favorecer proliferación microbiana y atracción de vectores.

Clasificación del riesgo: ALTO si no se corrige.
"""
            analisis_fotos.append(analisis_pro)

# -------------------------
# 2. EVALUACIÓN
# -------------------------
with tab2:
    diag_seg = st.selectbox("Segregación", ["CONFORME","PARCIAL","NO CONFORME"])
    diag_tra = st.selectbox("Trazabilidad", ["CONFORME","PARCIAL","NO CONFORME"])
    diag_equ = st.selectbox("Equipos", ["CONFORME","PARCIAL","NO CONFORME"])

# -------------------------
# 3. DIAGNÓSTICO MIP
# -------------------------
with tab3:
    riesgo = st.select_slider("Riesgo MIP", ["BAJO","MODERADO","ALTO","CRÍTICO"])
    plan = st.text_area("Plan de acción")

# -------------------------
# GENERAR INFORME DOCX
# -------------------------
def generar_docx():
    doc = Document()

    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = Pt(11)

    # TÍTULO
    doc.add_heading('INFORME TÉCNICO DE AUDITORÍA SANITARIA', 0)

    doc.add_paragraph(f"Cliente: {cliente}")
    doc.add_paragraph(f"Fecha: {fecha}")
    doc.add_paragraph(f"Auditor: {auditor}")
    doc.add_paragraph("Normativa: Resolución 2674 de 2013")

    # SECCIÓN 1
    doc.add_heading('1. Evidencia Fotográfica', 1)
    for a in analisis_fotos:
        doc.add_paragraph(a)

    # SECCIÓN 2
    doc.add_heading('2. Evaluación Técnica', 1)
    doc.add_paragraph(f"Segregación: {diag_seg}")
    doc.add_paragraph(f"Trazabilidad: {diag_tra}")
    doc.add_paragraph(f"Equipos: {diag_equ}")

    # SECCIÓN 3
    doc.add_heading('3. Diagnóstico MIP', 1)
    doc.add_paragraph(f"Nivel de riesgo: {riesgo}")

    # SECCIÓN 4
    doc.add_heading('4. Plan de Acción', 1)
    doc.add_paragraph(plan)

    # CONCLUSIÓN AUTOMÁTICA
    doc.add_heading('5. Conclusión Profesional', 1)
    doc.add_paragraph("""
El establecimiento presenta condiciones que requieren intervención técnica inmediata.
Se recomienda implementar un programa estructurado de Manejo Integral Ambiental (MIA),
fortaleciendo controles preventivos, trazabilidad y saneamiento para mitigar riesgos sanitarios.
""")

    return doc

# BOTÓN
if st.button("🚀 Generar Informe Profesional"):
    doc = generar_docx()

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📥 Descargar Informe (.DOCX)",
        data=buffer,
        file_name=f"Informe_Zodion_{cliente}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



